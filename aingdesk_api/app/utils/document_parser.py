"""
文档解析器模块
参考electron项目的文档解析实现，支持多种文件格式
"""

import os
import zipfile
import xml.etree.ElementTree as ET
import re
from typing import Optional, List, Dict, Any
from pathlib import Path
import json
import pandas as pd
from docx import Document as DocxDocument
from docx.shared import Inches
import openpyxl
from openpyxl import load_workbook
import fitz  # PyMuPDF
import jieba


class DocumentParser:
    """文档解析器类，支持多种文件格式"""
    
    # 支持的文件类型映射
    FILE_TYPE_MAP = {
        # 文档类型
        '.docx': 'parse_docx',
        '.doc': 'parse_doc',
        
        # 表格类型
        '.xlsx': 'parse_xlsx',
        '.xls': 'parse_xls',
        '.csv': 'parse_csv',
        
        # 演示文稿类型
        '.pptx': 'parse_pptx',
        '.ppt': 'parse_ppt',
        
        # PDF文件
        '.pdf': 'parse_pdf',
        
        # 网页文件
        '.html': 'parse_html',
        '.htm': 'parse_html',
        
        # URL地址
        'http': 'parse_url',
        'https': 'parse_url',
        
        # 图片类型
        '.jpg': 'parse_image',
        '.jpeg': 'parse_image',
        '.png': 'parse_image',
        '.gif': 'parse_image',
        '.bmp': 'parse_image',
        '.webp': 'parse_image',
        '.ppm': 'parse_image',
        '.tiff': 'parse_image',
        
        # Markdown和文本文件
        '.md': 'parse_txt',
        '.markdown': 'parse_txt',
        '.txt': 'parse_txt',
        '.log': 'parse_txt',
        '.text': 'parse_txt',
        '.conf': 'parse_txt',
        '.cfg': 'parse_txt',
        '.ini': 'parse_txt',
        '.json': 'parse_txt',
    }
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        """获取文件扩展名"""
        if filename.startswith('https://'):
            return 'https'
        elif filename.startswith('http://'):
            return 'http'
        return os.path.splitext(filename)[1].lower()
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """检查文件是否受支持"""
        ext = DocumentParser.get_file_extension(filename)
        return ext in DocumentParser.FILE_TYPE_MAP
    
    @staticmethod
    def get_supported_extensions() -> List[str]:
        """获取所有支持的文件扩展名"""
        return list(DocumentParser.FILE_TYPE_MAP.keys())
    
    @staticmethod
    def parse_document(filename: str) -> Dict[str, Any]:
        """
        解析文档
        
        Args:
            filename: 文件路径
            
        Returns:
            Dict包含解析结果，格式：{
                'success': bool,
                'content': str,
                'error': str (可选)
            }
        """
        try:
            ext = DocumentParser.get_file_extension(filename)
            if ext not in DocumentParser.FILE_TYPE_MAP:
                return {
                    'success': False,
                    'content': '',
                    'error': f'不支持的文件格式: {ext}'
                }
            
            # 检查文件是否存在
            if not os.path.exists(filename):
                return {
                    'success': False,
                    'content': '',
                    'error': f'文件不存在: {filename}'
                }
            
            # 获取解析方法
            parse_method = getattr(DocumentParser, DocumentParser.FILE_TYPE_MAP[ext])
            
            # 执行解析
            content = parse_method(filename)
            
            if not content or content.strip() == '':
                return {
                    'success': False,
                    'content': '',
                    'error': '文档内容为空'
                }
            
            return {
                'success': True,
                'content': content
            }
            
        except Exception as e:
            return {
                'success': False,
                'content': '',
                'error': f'解析文档失败: {str(e)}'
            }
    
    @staticmethod
    def parse_docx(filename: str) -> str:
        """解析Word文档(.docx)"""
        try:
            doc = DocxDocument(filename)
            content = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # 提取表格内容
            for table in doc.tables:
                content.append("\n【表格】")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
                content.append("【表格结束】\n")
            
            return "\n".join(content)
            
        except Exception as e:
            raise Exception(f"解析Word文档失败: {str(e)}")
    
    @staticmethod
    def parse_doc(filename: str) -> str:
        """解析旧版Word文档(.doc) - 简化实现"""
        # .doc文件需要更复杂的解析，这里先返回提示
        # 实际项目中可以使用win32com或antiword等库
        return f"# {os.path.basename(filename)}\n\n[注意：.doc格式文档解析功能开发中，建议使用.docx格式]"
    
    @staticmethod
    def parse_xlsx(filename: str) -> str:
        """解析Excel文档(.xlsx)"""
        try:
            # 使用pandas读取Excel文件
            df_dict = pd.read_excel(filename, sheet_name=None, engine='openpyxl')
            
            content_parts = []
            content_parts.append(f"# {os.path.basename(filename)}\n")
            content_parts.append(f"此Excel文档包含 {len(df_dict)} 个工作表：\n")
            
            for sheet_name, df in df_dict.items():
                content_parts.append(f"\n## 工作表: {sheet_name}\n")
                
                if df.empty:
                    content_parts.append("*此工作表为空*\n")
                    continue
                
                # 生成Markdown表格
                headers = df.columns.tolist()
                content_parts.append('| ' + ' | '.join(str(h) for h in headers) + ' |')
                content_parts.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
                
                # 添加数据行
                for _, row in df.iterrows():
                    cells = [str(val) if pd.notna(val) else '' for val in row.values]
                    content_parts.append('| ' + ' | '.join(cells) + ' |')
                
                content_parts.append(f"\n*共 {len(df)} 行数据*\n")
            
            return "\n".join(content_parts)
            
        except Exception as e:
            raise Exception(f"解析Excel文档失败: {str(e)}")
    
    @staticmethod
    def parse_xls(filename: str) -> str:
        """解析旧版Excel文档(.xls)"""
        try:
            # 使用pandas读取.xls文件
            df_dict = pd.read_excel(filename, sheet_name=None, engine='xlrd')
            
            content_parts = []
            content_parts.append(f"# {os.path.basename(filename)}\n")
            content_parts.append(f"此Excel文档包含 {len(df_dict)} 个工作表：\n")
            
            for sheet_name, df in df_dict.items():
                content_parts.append(f"\n## 工作表: {sheet_name}\n")
                
                if df.empty:
                    content_parts.append("*此工作表为空*\n")
                    continue
                
                # 生成Markdown表格
                headers = df.columns.tolist()
                content_parts.append('| ' + ' | '.join(str(h) for h in headers) + ' |')
                content_parts.append('| ' + ' | '.join(['---'] * len(headers)) + ' |')
                
                # 添加数据行
                for _, row in df.iterrows():
                    cells = [str(val) if pd.notna(val) else '' for val in row.values]
                    content_parts.append('| ' + ' | '.join(cells) + ' |')
                
                content_parts.append(f"\n*共 {len(df)} 行数据*\n")
            
            return "\n".join(content_parts)
            
        except Exception as e:
            # 如果xlrd引擎失败，尝试使用openpyxl
            try:
                return DocumentParser.parse_xlsx(filename)
            except:
                raise Exception(f"解析.xls文档失败: {str(e)}")
    
    @staticmethod
    def parse_csv(file_path: str) -> str:
        """解析CSV文件"""
        try:
            # 检测文件编码，参考electron项目实现
            import chardet
            
            # 首先尝试检测BOM标记
            with open(file_path, 'rb') as f:
                buffer = f.read(4096)  # 读取前4KB用于检测
            
            # 检测BOM
            encoding = None
            if len(buffer) >= 3 and buffer[0] == 0xEF and buffer[1] == 0xBB and buffer[2] == 0xBF:
                encoding = 'utf-8-sig'
            else:
                # 使用chardet检测编码
                try:
                    result = chardet.detect(buffer)
                    if result and result['encoding']:
                        encoding = result['encoding']
                except Exception:
                    pass
            
            # 如果检测失败，尝试常见编码
            if not encoding:
                encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'latin1', 'cp1252']
                content = None
                
                for test_encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=test_encoding) as f:
                            content = f.read()
                        # 检查是否有乱码字符
                        if '�' not in content and '锟�' not in content and '鈩�' not in content:
                            encoding = test_encoding
                            break
                    except (UnicodeDecodeError, UnicodeError):
                        continue
                
                if not encoding:
                    # 如果所有编码都失败，使用二进制模式并忽略错误
                    with open(file_path, 'rb') as f:
                        content = f.read().decode('utf-8', errors='ignore')
                    encoding = 'utf-8'
            
            # 读取完整文件内容
            if encoding != 'utf-8':  # 如果检测到的不是utf-8，重新读取
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
            elif 'content' not in locals():  # 如果content还没有被设置
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            
            # 解析CSV内容，处理引号内的换行和逗号
            lines = content.split('\n')
            records = []
            headers = []
            current_row = []
            current_value = ''
            in_quotes = False
            
            for line_num, line in enumerate(lines):
                i = 0
                while i < len(line):
                    char = line[i]
                    next_char = line[i + 1] if i + 1 < len(line) else ''
                    
                    # 处理引号
                    if char == '"':
                        if not in_quotes:
                            in_quotes = True
                        elif next_char == '"':
                            # 转义的引号
                            current_value += '"'
                            i += 1  # 跳过下一个引号
                        else:
                            # 结束引号
                            in_quotes = False
                    # 处理分隔符
                    elif char == ',' and not in_quotes:
                        current_row.append(current_value)
                        current_value = ''
                    # 处理行结束
                    elif char == '\r' or (char == '\n' and not in_quotes):
                        if char == '\r' and next_char == '\n':
                            i += 1  # 跳过\n
                        if not in_quotes:
                            # 结束当前行
                            current_row.append(current_value)
                            if current_row or line_num == 0:  # 保留标题行
                                records.append(current_row)
                            current_row = []
                            current_value = ''
                            break
                        else:
                            current_value += char
                    # 其他字符
                    else:
                        current_value += char
                    
                    i += 1
                
                # 处理行尾
                if current_value or current_row:
                    if in_quotes:
                        current_value += '\n'  # 引号内的换行保留
                    else:
                        current_row.append(current_value)
                        records.append(current_row)
                        current_row = []
                        current_value = ''
            
            # 处理最后一行
            if current_value or current_row:
                current_row.append(current_value)
                records.append(current_row)
            
            # 如果没有记录，返回空结果
            if not records:
                return f"# CSV文件: {os.path.basename(file_path)}\n\n*CSV文件为空或格式不正确*"
            
            # 获取表头并清理BOM标记
            headers = [h.replace('\uFEFF', '').strip() for h in records[0]]
            
            # 过滤空行
            valid_records = []
            empty_rows = 0
            
            for i, row in enumerate(records[1:], 1):  # 跳过标题行
                # 跳过完全空白的行
                if not any(cell.strip() for cell in row):
                    empty_rows += 1
                    continue
                
                # 确保每行有正确的列数
                row_obj = {}
                for j, header in enumerate(headers):
                    if header:  # 只处理非空列名
                        value = row[j] if j < len(row) else ""
                        row_obj[header] = value
                
                if row_obj:  # 只添加非空行
                    valid_records.append(row_obj)
            
            # 生成统计信息
            total_rows = len(valid_records)
            total_cols = len([h for h in headers if h])  # 非空列数
            
            # 生成唯一值统计
            unique_values = {}
            for header in headers:
                if header:
                    unique_values[header] = set()
            
            for record in valid_records:
                for header, value in record.items():
                    if value and str(value).strip():
                        unique_values[header].add(str(value).strip())
            
            # 生成Markdown表格
            markdown_lines = []
            
            # 文件标题
            markdown_lines.append(f"# CSV文件: {os.path.basename(file_path)}")
            markdown_lines.append("")
            
            # 文件摘要
            markdown_lines.append("## 文件摘要")
            markdown_lines.append("")
            markdown_lines.append(f"- **文件名**: {os.path.basename(file_path)}")
            markdown_lines.append(f"- **总行数**: {total_rows}")
            markdown_lines.append(f"- **总列数**: {total_cols}")
            markdown_lines.append(f"- **空行数**: {empty_rows}")
            markdown_lines.append("")
            
            # 列信息
            if total_cols > 0:
                markdown_lines.append("### 列信息")
                markdown_lines.append("")
                markdown_lines.append("| 列名 | 唯一值数量 | 样例值 |")
                markdown_lines.append("| --- | --- | --- |")
                
                for header in headers:
                    if header:  # 只显示非空列名
                        unique_count = len(unique_values.get(header, set()))
                        sample_values = list(unique_values.get(header, set()))[:3]
                        sample_str = ", ".join(str(v)[:30] for v in sample_values)
                        
                        # 转义Markdown特殊字符
                        header_escaped = header.replace("|", "\\|")
                        sample_escaped = sample_str.replace("|", "\\|")
                        
                        markdown_lines.append(f"| {header_escaped} | {unique_count} | {sample_escaped} |")
                
                markdown_lines.append("")
            
            # 数据表格
            if total_rows > 0 and total_cols > 0:
                markdown_lines.append("### 数据表格")
                markdown_lines.append("")
                
                # 表头
                header_row = "| " + " | ".join(h.replace("|", "\\|") for h in headers if h) + " |"
                markdown_lines.append(header_row)
                
                # 分隔符
                separator_row = "| " + " | ".join("---" for h in headers if h) + " |"
                markdown_lines.append(separator_row)
                
                # 数据行（限制显示前100行）
                display_rows = min(100, len(valid_records))
                for i in range(display_rows):
                    record = valid_records[i]
                    row_data = []
                    for header in headers:
                        if header and header in record:
                            value = str(record[header]).replace("|", "\\|")
                            row_data.append(value)
                        elif header:
                            row_data.append("")
                    
                    row_str = "| " + " | ".join(row_data) + " |"
                    markdown_lines.append(row_str)
                
                if len(valid_records) > 100:
                    markdown_lines.append("")
                    markdown_lines.append(f"*显示前100行，共{len(valid_records)}行*")
            
            return "\n".join(markdown_lines)
            
        except Exception as e:
            print(f"CSV解析失败: {e}")
            return f"# CSV解析失败\n\n无法解析CSV文件。错误: {str(e)}"
    
    @staticmethod
    def parse_pptx(filename: str) -> str:
        """解析PowerPoint文档(.pptx)"""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            # 读取PPTX文件（实际上是一个ZIP文件）
            with zipfile.ZipFile(filename, 'r') as zip_file:
                # 获取presentation.xml文件内容
                presentation_xml = None
                try:
                    presentation_xml = zip_file.read('ppt/presentation.xml').decode('utf-8')
                except KeyError:
                    return f"# {os.path.basename(filename)}\n\n[无效的PPTX文件：缺少presentation.xml]"
                
                # 提取幻灯片ID
                slide_ids = []
                if presentation_xml:
                    slide_count_match = re.search(r'<p:sldIdLst>([^]*?)</p:sldIdLst>', presentation_xml)
                    if slide_count_match:
                        slide_ids = re.findall(r'id="(\d+)"', slide_count_match.group(1))
                
                content_parts = []
                content_parts.append(f"# {os.path.basename(filename)}\n")
                content_parts.append(f"此PowerPoint文档包含 {len(slide_ids)} 张幻灯片：\n")
                
                # 处理每个幻灯片
                for i, slide_id in enumerate(slide_ids, 1):
                    try:
                        slide_xml = zip_file.read(f'ppt/slides/slide{i}.xml').decode('utf-8')
                        
                        # 提取文本内容
                        paragraphs = []
                        paragraph_elements = re.findall(r'<a:p>.*?</a:p>', slide_xml)
                        
                        for paragraph in paragraph_elements:
                            text_elements = re.findall(r'<a:t>(.+?)</a:t>', paragraph)
                            if text_elements:
                                paragraph_text = ' '.join(text_elements)
                                if paragraph_text.strip():
                                    paragraphs.append(paragraph_text.strip())
                        
                        if paragraphs:
                            content_parts.append(f"\n## 幻灯片 {i}\n")
                            content_parts.append('\n'.join(paragraphs))
                            
                    except KeyError:
                        content_parts.append(f"\n## 幻灯片 {i}\n[无法读取幻灯片内容]")
                
                return '\n'.join(content_parts)
                
        except Exception as e:
            return f"# {os.path.basename(filename)}\n\n[PowerPoint文档解析失败: {str(e)}]"
    
    @staticmethod
    def parse_ppt(filename: str) -> str:
        """解析旧版PowerPoint文档(.ppt)"""
        return f"# {os.path.basename(filename)}\n\n[注意：目前仅支持.pptx格式的PowerPoint文件解析]"
    
    @staticmethod
    def parse_pdf(filename: str) -> str:
        """解析PDF文档"""
        try:
            doc = fitz.open(filename)
            content = []
            
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    content.append(text.strip())
            
            doc.close()
            
            # 如果文本提取失败，尝试OCR
            if not content or not any(c.strip() for c in content):
                return DocumentParser.parse_pdf_with_ocr(filename)
            
            return "\n".join(content)
            
        except Exception as e:
            raise Exception(f"解析PDF文档失败: {str(e)}")
    
    @staticmethod
    def parse_pdf_with_ocr(filename: str) -> str:
        """使用OCR解析PDF文档"""
        try:
            import pytesseract
            from PIL import Image
            import io
            
            doc = fitz.open(filename)
            content = []
            
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                # 将页面转换为图片
                pix = page.get_pixmap()
                img_data = pix.tobytes("ppm")
                
                # 使用OCR识别文本
                try:
                    img = Image.open(io.BytesIO(img_data))
                    text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                    if text.strip():
                        content.append(text.strip())
                except Exception as ocr_error:
                    print(f"OCR处理页面 {page_num + 1} 失败: {ocr_error}")
                
                pix = None
            
            doc.close()
            
            if content:
                return "\n".join(content)
            else:
                return f"# {os.path.basename(filename)}\n\n[PDF文档内容提取失败，可能需要OCR支持]"
                
        except ImportError:
            return f"# {os.path.basename(filename)}\n\n[PDF文档需要OCR支持，请安装pytesseract]"
        except Exception as e:
            return f"# {os.path.basename(filename)}\n\n[PDF文档OCR处理失败: {str(e)}]"
    
    @staticmethod
    def parse_html(file_path: str) -> str:
        """解析HTML文件或URL"""
        try:
            # 检查是否为URL
            if file_path.startswith('http://') or file_path.startswith('https://'):
                return DocumentParser.parse_url(file_path)
            
            # 本地HTML文件
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # 使用正则表达式进行HTML解析，参考electron项目实现
            # 移除干扰元素
            interference_patterns = [
                r'<script.*?</script>',  # script标签
                r'<style.*?</style>',     # style标签
                r'<nav.*?</nav>',         # nav标签
                r'<footer.*?</footer>',   # footer标签
                r'<aside.*?</aside>',     # aside标签
                r'<header.*?</header>',   # header标签
                r'<div[^>]*class[^>]*["\'][^"\']*(?:Header|Footer|Sidebar|Ads|Banner|Advertisement|Copyright)[^"\']*["\'][^>]*>.*?</div>',  # 特定class的div
                r'<[^>]*class[^>]*["\'][^"\']*(?:Header|Footer|Sidebar|Ads|Banner|Advertisement|Copyright)[^"\']*["\'][^>]*>',  # 特定class的任何标签
            ]
            
            for pattern in interference_patterns:
                content = re.sub(pattern, '', content, flags=re.DOTALL | re.IGNORECASE)
            
            # 提取标题
            title_match = re.search(r'<title[^>]*>(.*?)</title>', content, re.IGNORECASE | re.DOTALL)
            title = title_match.group(1).strip() if title_match else ''
            
            # 查找主要内容区域
            main_content = ''
            
            # 优先查找特定标签
            content_selectors = [
                r'<article[^>]*>(.*?)</article>',
                r'<div[^>]*class[^>]*["\'][^"\']*article[^"\']*["\'][^>]*>(.*?)</div>',
                r'<div[^>]*id[^>]*["\'][^"\']*article[^"\']*["\'][^>]*>(.*?)</div>',
                r'<div[^>]*class[^>]*["\'][^"\']*content_text[^"\']*["\'][^>]*>(.*?)</div>',
                r'<div[^>]*id[^>]*["\'][^"\']*content_text[^"\']*["\'][^>]*>(.*?)</div>',
            ]
            
            for selector in content_selectors:
                match = re.search(selector, content, re.IGNORECASE | re.DOTALL)
                if match:
                    main_content = match.group(1)
                    break
            
            # 如果没有找到主要内容，使用整个内容
            if not main_content:
                main_content = content
            
            # 提取文本内容
            text_content = ''
            
            # 提取标题
            if title:
                text_content += f"# {title}\n\n"
            
            # 提取段落
            paragraph_matches = re.findall(r'<p[^>]*>(.*?)</p>', main_content, re.IGNORECASE | re.DOTALL)
            for p in paragraph_matches:
                text = re.sub(r'<[^>]+>', '', p).strip()
                if text and len(text) > 10:  # 过滤短文本
                    text_content += f"{text}\n\n"
            
            # 提取没有块级元素的div中的文本
            div_matches = re.findall(r'<div[^>]*>(.*?)</div>', main_content, re.IGNORECASE | re.DOTALL)
            for div in div_matches:
                # 检查是否包含块级元素
                if not re.search(r'<(p|h[1-6]|ul|ol|table|blockquote)', div, re.IGNORECASE):
                    text = re.sub(r'<[^>]+>', '', div).strip()
                    if text and len(text) > 10:
                        text_content += f"{text}\n\n"
            
            # 提取标题
            for i in range(1, 7):
                h_matches = re.findall(rf'<h{i}[^>]*>(.*?)</h{i}>', main_content, re.IGNORECASE | re.DOTALL)
                for h in h_matches:
                    text = re.sub(r'<[^>]+>', '', h).strip()
                    if text:
                        text_content += f"{'#' * i} {text}\n\n"
            
            # 提取列表
            list_matches = re.findall(r'<(ul|ol)[^>]*>(.*?)</\1>', main_content, re.IGNORECASE | re.DOTALL)
            for list_type, list_content in list_matches:
                li_matches = re.findall(r'<li[^>]*>(.*?)</li>', list_content, re.IGNORECASE | re.DOTALL)
                for idx, li in enumerate(li_matches):
                    text = re.sub(r'<[^>]+>', '', li).strip()
                    if text:
                        if list_type == 'ul':
                            text_content += f"* {text}\n"
                        else:
                            text_content += f"{idx + 1}. {text}\n"
                text_content += "\n"
            
            # 提取表格
            table_matches = re.findall(r'<table[^>]*>(.*?)</table>', main_content, re.IGNORECASE | re.DOTALL)
            for table_content in table_matches:
                # 提取表头
                thead_matches = re.findall(r'<thead[^>]*>(.*?)</thead>', table_content, re.IGNORECASE | re.DOTALL)
                headers = []
                for thead in thead_matches:
                    th_matches = re.findall(r'<th[^>]*>(.*?)</th>', thead, re.IGNORECASE | re.DOTALL)
                    for th in th_matches:
                        text = re.sub(r'<[^>]+>', '', th).strip()
                        if text:
                            headers.append(text)
                
                # 提取表体
                tbody_matches = re.findall(r'<tbody[^>]*>(.*?)</tbody>', table_content, re.IGNORECASE | re.DOTALL)
                if not tbody_matches:
                    tbody_matches = [table_content]  # 如果没有tbody，直接使用整个表格内容
                
                rows = []
                for tbody in tbody_matches:
                    tr_matches = re.findall(r'<tr[^>]*>(.*?)</tr>', tbody, re.IGNORECASE | re.DOTALL)
                    for tr in tr_matches:
                        td_matches = re.findall(r'<td[^>]*>(.*?)</td>', tr, re.IGNORECASE | re.DOTALL)
                        row = []
                        for td in td_matches:
                            text = re.sub(r'<[^>]+>', '', td).strip()
                            row.append(text)
                        if row and any(row):  # 确保行不为空
                            rows.append(row)
                
                # 如果没有提取到表头，尝试从第一行获取
                if not headers and rows:
                    headers = rows[0]
                    rows = rows[1:]
                
                # 生成Markdown表格
                if headers or rows:
                    text_content += "\n"
                    if headers:
                        text_content += "| " + " | ".join(headers) + " |\n"
                        text_content += "| " + " | ".join(["---"] * len(headers)) + " |\n"
                    
                    for row in rows:
                        if len(row) == len(headers) or not headers:
                            text_content += "| " + " | ".join(row) + " |\n"
                    
                    text_content += "\n"
            
            # 清理多余的空白字符
            text_content = re.sub(r'\n{3,}', '\n\n', text_content)
            text_content = re.sub(r'[ \t]+', ' ', text_content)
            
            return text_content.strip()
        except Exception as e:
            return ""
    
    @staticmethod
    def parse_image(filename: str) -> str:
        """解析图片文件 - 支持OCR文字识别"""
        try:
            # 首先尝试OCR识别
            ocr_result = DocumentParser.parse_image_with_ocr(filename)
            if ocr_result and ocr_result.strip():
                return ocr_result
            
            # 如果OCR失败，返回基本信息
            from PIL import Image
            
            with Image.open(filename) as img:
                width, height = img.size
                format_name = img.format
                
            return f"# {os.path.basename(filename)}\n\n图片信息:\n- 尺寸: {width} x {height} 像素\n- 格式: {format_name}"
            
        except Exception as e:
            return f"# {os.path.basename(filename)}\n\n图片文件，无法提取文本内容"
    
    @staticmethod
    def parse_image_with_ocr(filename: str) -> str:
        """使用OCR解析图片中的文字"""
        try:
            import pytesseract
            from PIL import Image
            
            # 打开图片
            with Image.open(filename) as img:
                # 使用OCR识别文字
                text = pytesseract.image_to_string(img, lang='chi_sim+eng')
                
            if text and text.strip():
                # 清理识别结果
                text = text.replace('\n\n\n', '\n\n').strip()
                return f"# {os.path.basename(filename)}\n\n识别的文字内容:\n\n{text}"
            else:
                return ""
                
        except ImportError:
            # pytesseract未安装，返回空字符串让外层处理
            return ""
        except Exception as e:
            return ""
    
    @staticmethod
    def parse_txt(file_path: str) -> str:
        """解析TXT文件"""
        try:
            # 尝试多种编码方式读取文件
            encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'iso-8859-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，尝试二进制读取并解码
            if content is None:
                try:
                    with open(file_path, 'rb') as f:
                        content = f.read().decode('utf-8', errors='ignore')
                except Exception:
                    return ""
            
            return content
        except Exception as e:
            return ""
    
    @staticmethod
    def parse_url(filename: str) -> str:
        """解析URL地址 - 返回URL信息"""
        try:
            return f"# URL地址\n\n地址: {filename}\n\n[这是一个URL地址，需要通过网络请求获取内容]"
        except Exception as e:
            raise Exception(f"解析URL失败: {str(e)}")


def parse_document(filename: str, rag_name: str = '') -> Dict[str, Any]:
    """
    解析文档的主函数
    
    Args:
        filename: 文件路径
        rag_name: 知识库名称（可选）
        
    Returns:
        解析结果字典
    """
    return DocumentParser.parse_document(filename)


def is_supported_file_type(filename: str) -> bool:
    """检查文件类型是否受支持"""
    return DocumentParser.is_supported(filename)


def get_supported_file_extensions() -> List[str]:
    """获取所有支持的文件扩展名"""
    return DocumentParser.get_supported_extensions()